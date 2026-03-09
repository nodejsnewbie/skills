import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, font_name):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def create_exam_paper():
    doc = Document()
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2024-2025 学年第一学期 《Java高级编程》 期末试卷（A卷）")
    run.font.size = Pt(16)
    run.bold = True
    set_font(run, '黑体')

    # 一、单项选择题
    p = doc.add_paragraph()
    run = p.add_run("一、单项选择题（每题 2 分，共 30 分）")
    run.font.size = Pt(14)
    run.bold = True
    set_font(run, '黑体')
    
    questions = [
        "1. [第一章] 以下不属于软件设计模式基本原则的是？（ D ）\nA. 单一职责原则  B. 开闭原则  C. 依赖混淆原则  D. 迪米特法则",
        "2. [第一章] 设计模式的核心思想是？（ B ）\nA. 提高代码执行速度  B. 提高代码的可复用性、可读性和可维护性  C. 减少代码行数  D. 避免报错",
        "3. [第二章] 保证一个类只有一个实例，并提供一个访问它的全局访问点的模式是？（ B ）\nA. 工厂模式 B. 单例模式 C. 建造者模式 D. 原型模式",
        "4. [第二章] 哪种单例模式在类加载时就完成了初始化？（ B ）\nA. 懒汉式 B. 饿汉式 C. 静态内部类 D. 枚举类",
        "5. [第三章] 抽象工厂模式的主要目的是？（ B ）\nA. 创建单一对象 B. 创建一系列相关或相互依赖的对象 C. 销毁对象 D. 代理对象",
        "6. [第三章] 在抽象工厂模式中，增加一个新的产品族（比如全新的操作系统），正确的是？（ A ）\nA. 只需要增加具体工厂即可 B. 需要修改抽象类 C. 违反开闭原则 D. 违反单一职责",
        "7. [第四章] 将一个类的接口转换成客户希望的另外一个接口。这是哪种模式？（ B ）\nA. 代理模式 B. 适配器模式 C. 桥接模式 D. 外观模式",
        "8. [第四章] 适配器模式主要分为哪两种实现方式？（ B ）\nA. 接口与抽象 B. 类与对象 C. 组件与对象 D. 桥接与组合",
        "9. [第五章] 在不改变对象结构的情况下，动态给对象添加额外职责的是？（ B ）\nA. 桥接模式 B. 装饰模式 C. 代理模式 D. 外观模式",
        "10. [第五章] 组合模式用于表示对象的？（ C ）\nA. 网状结构 B. 线性结构 C. 部分-整体的树形结构 D. 无依赖结构",
        "11. [第六章] 定义算法族，分别封装起来，让它们之间互换。这是指？（ B ）\nA. 模板方法 B. 策略模式 C. 状态模式 D. 观察者模式",
        "12. [第六章] 在模板方法模式中，定义算法骨架的方法通常声明为？（ C ）\nA. abstract B. private C. final D. public",
        "13. [第七章] 状态模式中，具体状态相关的行为通常封装在哪里？（ C ）\nA. Context类 B. State接口 C. ConcreteState类 D. Client端",
        "14. [第七章] 观察者模式通常用于实现什么关系的通信？（ B ）\nA. 一对一 B. 一对多 C. 多对一 D. 多对多",
        "15. [第八章] 访问者模式的核心是？（ C ）\nA. 隐藏调用 B. 改变状态 C. 在不修改现有类的前提下增加新操作 D. 控制访问"
    ]
    for q in questions:
        doc.add_paragraph(q)

    # 二、填空题
    p = doc.add_paragraph()
    run = p.add_run("二、填空题（每空 2 分，共 10 分）")
    run.font.size = Pt(14)
    run.bold = True
    set_font(run, '黑体')
    blanks = [
        "16. [第一章] 设计模式的开闭原则是指软件实体应对__扩展__开放，对__修改__关闭。",
        "17. [第二章] 单例模式懒汉式中为保证线程安全常使用__synchronized__关键字。",
        "18. [第二章] 原型模式通过重写 Java 中 Object 类的__clone()__方法实现克隆。",
        "19. [第五章] 外观模式的作用是为子系统的接口提供一个统一的__高层接口/访问接口__。",
        "20. [第七章] 提供顺序访问聚合对象元素而不暴露内部结构的模式是__迭代器__模式。"
    ]
    for q in blanks:
        doc.add_paragraph(q.replace('__扩展__', '______')
                           .replace('__修改__', '______')
                           .replace('__synchronized__', '______')
                           .replace('__clone()__', '______')
                           .replace('__高层接口/访问接口__', '______')
                           .replace('__迭代器__', '______'))

    # 三、判断题
    p = doc.add_paragraph()
    run = p.add_run("三、判断题（每题 1 分，共 10 分）")
    run.font.size = Pt(14)
    run.bold = True
    set_font(run, '黑体')
    tfs = [
        "21. [第一章] 依赖倒置原则要求针对接口编程，不要针对实现编程。（ ）",
        "22. [第二章] 饿汉式单例模式在类加载时就完成初始化，不存在多线程安全问题。（ ）",
        "23. [第二章] 工厂方法模式扩展产品只需增加新工厂类即可，原代码无需修改。（ ）",
        "24. [第二章] 原型模式绝对无法实现深克隆。（ ）",
        "25. [第四章] Java 不支持多重继承，因此类适配器通过继承类和实现接口完成。（ ）",
        "26. [第五章] 装饰模式过渡使用会增加许多小类，增加系统复杂性。（ ）",
        "27. [第六章] 策略模式中，客户端不需要知道各种策略的具体实现。（ ）",
        "28. [第六章] 模板方法模式的骨架由具体子类定义。（ ）",
        "29. [第七章] 观察者模式中，主题和观察者是低耦合的。（ ）",
        "30. [第八章] 访问者模式非常适合对象结构频繁变动的系统。（ ）"
    ]
    for q in tfs:
        doc.add_paragraph(q)

    # 四、简答题
    p = doc.add_paragraph()
    run = p.add_run("四、简答题（每题 6 分，共 30 分）")
    run.font.size = Pt(14)
    run.bold = True
    set_font(run, '黑体')
    sas = [
        "31. [第一章] 简述单一职责原则（SRP）的基本定义及意义。",
        "32. [第二章] 比较简单工厂模式与工厂方法模式的区别。",
        "33. [第四章] 简述桥接模式的优缺点及核心意图。",
        "34. [第六章] 简述策略模式的优缺点，并在什么场景下适合使用？",
        "35. [第八章] 请简述访问者模式的适用场景及其弊端。"
    ]
    for q in sas:
        doc.add_paragraph(q)
        doc.add_paragraph("\n")

    # 五、综合题
    p = doc.add_paragraph()
    run = p.add_run("五、综合分析与设计题（共 20 分）")
    run.font.size = Pt(14)
    run.bold = True
    set_font(run, '黑体')
    doc.add_paragraph("36. [第三章] (10分)\n某手机公司的跨平台组件系统同时生产 BrandA 和 BrandB 的手机，具有 Screen 和 Battery 部件。\n请使用抽象工厂模式设计该系统：说明包含的抽象/具体角色（4分），并简要给出类结构关系或代码声明（6分）。\n\n\n\n\n")
    doc.add_paragraph("37. [第八章] (10分)\n人力资源系统中，有 FullTimeEmployee 和 PartTimeEmployee。现需增加计算工资发薪和统计休假两项操作。\n请运用访问者模式设计该功能：指出 Visitor 接口/实现及 Element 接口/实现（5分），并说明此模式在此场景的设计优势（5分）。\n\n\n\n\n")

    doc.save('1 试卷纸-A卷.docx')

def create_answer_sheet():
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2024-2025 学年第一学期 《Java高级编程》 期末试卷（A卷）答题卡")
    run.font.size = Pt(16)
    run.bold = True
    set_font(run, '黑体')

    doc.add_paragraph("学号：_______________  姓名：_______________  班级：_______________\n")

    doc.add_paragraph("一、单项选择题（每题 2 分，共 30 分）")
    table = doc.add_table(rows=4, cols=8)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i in range(8): hdr_cells[i].text = f"题号" if i==0 else str(i)
    ans_cells = table.rows[1].cells
    ans_cells[0].text = "答案"
    hdr_cells2 = table.rows[2].cells
    for i in range(8): hdr_cells2[i].text = f"题号" if i==0 else str(i+7)
    ans_cells2 = table.rows[3].cells
    ans_cells2[0].text = "答案"

    doc.add_paragraph("\n二、填空题（每空 2 分，共 10 分）")
    doc.add_paragraph("16 ____________  ____________   17 ____________   18 ____________   19 ____________   20 ____________")

    doc.add_paragraph("\n三、判断题（每题 1 分，共 10 分）")
    table2 = doc.add_table(rows=2, cols=11)
    table2.style = 'Table Grid'
    h2 = table2.rows[0].cells
    h2[0].text = "题号"
    for i in range(1, 11): h2[i].text = str(i+20)
    a2 = table2.rows[1].cells
    a2[0].text = "答案"

    doc.add_paragraph("\n四、简答题（每题 6 分，共 30 分）")
    for i in range(31, 36):
        doc.add_paragraph(f"{i}.")
        doc.add_paragraph("\n\n")
    
    doc.add_paragraph("\n五、综合题（共 20 分）")
    doc.add_paragraph("36.")
    doc.add_paragraph("\n\n\n\n")
    doc.add_paragraph("37.")
    doc.add_paragraph("\n\n\n\n")

    doc.save('2 答题纸-A卷.docx')

def create_answer_key():
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2024-2025 学年第一学期 《Java高级编程》 期末试卷（A卷）参考答案及评分标准")
    run.font.size = Pt(16)
    run.bold = True
    set_font(run, '黑体')

    doc.add_paragraph("一、单项选择题（每题 2 分，共 30 分）\n1-5 D B B B B  6-10 A B B B C  11-15 B C C B C")
    doc.add_paragraph("二、填空题（每空 2 分，共 10 分）\n16 扩展 修改（各1分）\n17 synchronized\n18 clone()\n19 统一的高层接口访问\n20 迭代器")
    doc.add_paragraph("三、判断题（每题 1 分，共 10 分）\n21-25 √ √ √ × √   26-30 √ × × √ ×")
    doc.add_paragraph("四、简答题（每题 6 分，共 30 分）\n31. 一个类应该仅有一个引起它变化的原因。（3分）能够降低类的复杂度，提高可读性，降低变更引起的风险。（3分）\n32. 简单工厂决定由一个类创建实例，违背开闭原则；工厂方法将实例化延迟到子类，符合开闭原则。（各3分）\n33. 核心是将抽象部分与实现部分分离使之独立变化。（2分）优点提高了扩展性实现了接口复用（2分）；缺点增加了系统复杂度（2分）。\n34. 优点算法可自由切换、避免多重判断条件、扩展性良好；缺点客户端需知晓所有策略、类数量增加。（3分）；适用需动态在几种算法切换且行为独立变化的场景。（3分）\n35. 适用对象结构稳定但需要在此基础上定义新操作，分离数据与行为；由于新增元素需要修改所有Visitor类，违背开闭原则，因而不适合元素结构不稳定的场景。（各3分）")
    doc.add_paragraph("五、综合分析与设计题（共 20 分）\n36. （10分）\n角色包含：抽象工厂如PhoneFactory；具体工厂如BrandAFactory、BrandBFactory；抽象产品如Screen、Battery；具体产品如 BrandAScreen等。（4分）\n结构关系声明必须包含 `Screen createScreen()` 等抽象方法及继承实现。（6分）\n\n37. （10分）\n角色：Visitor如EmployeeDataVisitor，及其实现类SalaryCalculatorVisitor等；Element如Employee，及其实现类FullTimeEmployee等。（5分）\n优势：实现了数据结构与计算行为分离，若需增新行为（如核算绩效）只需加新访问者类，完全符合单一职责与开闭原则。（5分）")
    
    doc.save('3 参考答案及评分标准-A卷.docx')

def create_specification():
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2024-2025 学年第一学期 《Java高级编程》 命题双向细目表汇总")
    run.font.size = Pt(16)
    run.bold = True
    set_font(run, '黑体')

    table = doc.add_table(rows=7, cols=8)
    table.style = 'Table Grid'
    headers = ['章节名称', '题型', '认知(分)', '理解(分)', '应用(分)', '分析(分)', '综合(分)', '小计']
    for i, h in enumerate(headers): table.rows[0].cells[i].text = h
    
    data = [
        ['第一章', '选择,填空,判断,简答', '3', '4', '6', '0', '0', '13'],
        ['第二章', '选择,填空,判断,简答', '5', '5', '6', '0', '0', '16'],
        ['第三章', '选择,综合', '4', '0', '10', '0', '0', '14'],
        ['第四章', '选择,判断,简答', '4', '1', '6', '0', '0', '11'],
        ['第五章', '选择,填空,判断', '4', '2', '0', '0', '0', '6'],
        ['第六78章', '选择,填空,判断,简答,综合', '10', '8', '2', '0', '20', '40'],
    ]
    for r, row in enumerate(data):
        for c, val in enumerate(row):
            table.rows[r+1].cells[c].text = val
    
    doc.add_paragraph("\n题型总计分值：\n选择题: 30分\n填空题: 10分\n判断题: 10分\n简答题: 30分\n综合题: 20分\n共计100分。")
    doc.save('4 命题双向细目表汇总-A卷.docx')

if __name__ == '__main__':
    create_exam_paper()
    create_answer_sheet()
    create_answer_key()
    create_specification()
    print("FINISHED DOCX GENERATION")
